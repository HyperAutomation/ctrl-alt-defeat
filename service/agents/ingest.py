"""Component 1 — Ingest & Segmentation.

Non-LLM executor. Parses MD / TXT / PDF / DOCX into `Document` records with
page/paragraph anchors. The anchors are the traceability invariant — every
later citation must resolve to one of them (plan §4.2).
"""
from __future__ import annotations

import re
from pathlib import Path

try:  # pure helpers below must be importable without agent_framework
    from agent_framework import Executor, WorkflowContext, handler  # type: ignore
except ImportError:  # pragma: no cover -- tests for pure helpers skip MAF
    Executor = object  # type: ignore[assignment,misc]
    WorkflowContext = None  # type: ignore[assignment]

    def handler(fn):  # type: ignore[no-redef]
        return fn

from workflow.schemas import Document, DocumentSection, Submission


def _segment_markdown_or_text(text: str, filename: str) -> list[DocumentSection]:
    """Header-aware splitter. Works for .md and .txt; tolerates no headers."""
    sections: list[DocumentSection] = []
    # Simple heuristic: split on lines starting with '#' (markdown) or lines
    # that look like numbered section headings (e.g. "1.", "1.1").
    lines = text.splitlines()
    buf: list[str] = []
    current_heading: str | None = None
    current_page = 1
    paragraph_idx = 0

    def flush():
        nonlocal paragraph_idx, buf
        if not buf:
            return
        joined = "\n".join(buf).strip()
        if joined:
            paragraph_idx += 1
            sections.append(
                DocumentSection(
                    section_id=f"{filename}#{paragraph_idx}",
                    heading=current_heading,
                    page=current_page,
                    paragraph=paragraph_idx,
                    text=joined,
                )
            )
        buf = []

    page_marker = re.compile(r"<!--\s*Page\s+(\d+)\s*-->", re.IGNORECASE)
    md_heading = re.compile(r"^\s*#{1,6}\s+(.*)")
    numbered_heading = re.compile(r"^\s*\d+(\.\d+)*\s+(.+)")

    for line in lines:
        mp = page_marker.match(line)
        if mp:
            flush()
            current_page = int(mp.group(1))
            continue
        mh = md_heading.match(line)
        if mh:
            flush()
            current_heading = mh.group(1).strip()
            continue
        mn = numbered_heading.match(line)
        if mn:
            flush()
            current_heading = mn.group(2).strip()
            continue
        buf.append(line)
    flush()

    if not sections:
        # Fallback — whole document as one section.
        sections.append(
            DocumentSection(
                section_id=f"{filename}#1",
                heading=None,
                page=1,
                paragraph=1,
                text=text.strip(),
            )
        )
    return sections


def parse_file(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        text = path.read_text(errors="replace")
        sections = _segment_markdown_or_text(text, path.name)
        mime = "text/markdown" if suffix == ".md" else "text/plain"
    elif suffix == ".pdf":
        text = _parse_pdf(path)
        sections = _segment_markdown_or_text(text, path.name)
        mime = "application/pdf"
    elif suffix in {".docx"}:
        text = _parse_docx(path)
        sections = _segment_markdown_or_text(text, path.name)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return Document(
        document_id=path.name,
        filename=path.name,
        mime_type=mime,
        sections=sections,
    )


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover -- dev dependency missing
        return path.read_bytes().decode(errors="replace")
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        parts.append(f"<!-- Page {i} -->\n{page.extract_text() or ''}")
    return "\n".join(parts)


def _parse_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError:  # pragma: no cover
        return ""
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


class IngestExecutor(Executor):
    """Consumes a `Submission` whose `documents` list has file paths but no
    sections; returns the same `Submission` with `documents` fully segmented.

    The input contract uses the `Document.filename` as a path relative to the
    workspace root, which the FastAPI app makes available via ctx state.
    """

    def __init__(self) -> None:
        super().__init__("ingest")

    @handler
    async def ingest(self, submission: Submission, ctx: WorkflowContext[Submission]) -> None:
        workspace_root = None
        if hasattr(ctx, "get_state"):
            workspace_root = ctx.get_state("workspace_root")
        workspace_root = Path(workspace_root) if workspace_root else Path.cwd()

        parsed_docs: list[Document] = []
        for doc in submission.documents:
            if doc.sections:
                parsed_docs.append(doc)
                continue
            path = workspace_root / doc.filename
            if not path.exists():
                path = Path(doc.filename)  # absolute-path fallback
            parsed_docs.append(parse_file(path))

        submission.documents = parsed_docs
        await ctx.send_message(submission)
